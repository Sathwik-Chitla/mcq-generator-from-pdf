import streamlit as st
import fitz  # PyMuPDF
from langchain_ollama import ChatOllama
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# App Title
st.set_page_config(page_title="MCQ Generator", page_icon="📘", layout="wide")
st.title("📘 PDF → MCQ Generator")

# Upload PDF
uploaded_file = st.file_uploader("Upload a PDF", type="pdf")

# Number of MCQs input
num_q = st.number_input("How many MCQs do you want?", min_value=1, max_value=50, value=5, step=1)

if uploaded_file is not None:
    # Extract text from PDF
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = " ".join(page.get_text() for page in doc)

    # LLM + Prompt
    MODEL = "llama3.1"
    llm = ChatOllama(model=MODEL, temperature=0)

    prompt = PromptTemplate(
    input_variables=["content", "num_q"],
    template="""
    Generate {num_q} multiple choice questions (MCQs) from the following content.  
    Each question should have 4 options (A, B, C, D) with exactly one correct answer.  
    Format the output as JSON with this structure:
    [
    {{
      "question": "Question text",
      "options": ["A) ...", "B) ...", "C) ...", "D) ..."],
      "answer": "Correct Option"
    }},
    ...
    ]
    Content:
    {content}
    """
)

    parser = JsonOutputParser()
    chain = prompt | llm | parser

    if st.button("✨ Generate MCQs"):
        with st.spinner("Generating MCQs... ⏳"):
            mcqs = chain.invoke({"content": text[:4000], "num_q": num_q})

        st.success(f"✅ Generated {len(mcqs)} MCQs!")

        # Display MCQs
        for i, mcq in enumerate(mcqs, 1):
            st.markdown(f"**Q{i}. {mcq['question']}**")
            for opt in mcq["options"]:
                st.markdown(f"- {opt}")
            st.markdown(f"✅ **Answer:** {mcq['answer']}")
            st.write("---")

        # Save as Word
        def save_as_docx(mcqs):
            doc = Document()
            doc.add_heading("Generated MCQs", level=1)
            for i, mcq in enumerate(mcqs, 1):
                doc.add_paragraph(f"Q{i}. {mcq['question']}")
                for opt in mcq["options"]:
                    doc.add_paragraph(f"- {opt}", style="List Bullet")
                doc.add_paragraph(f"Answer: {mcq['answer']}\n")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        # Save as PDF
        def save_as_pdf(mcqs):
            buffer = BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            y = height - 50
            c.setFont("Helvetica-Bold", 14)
            c.drawString(200, y, "Generated MCQs")
            y -= 40
            c.setFont("Helvetica", 12)

            for i, mcq in enumerate(mcqs, 1):
                c.drawString(50, y, f"Q{i}. {mcq['question']}")
                y -= 20
                for opt in mcq["options"]:
                    c.drawString(70, y, f"- {opt}")
                    y -= 20
                c.drawString(50, y, f"Answer: {mcq['answer']}")
                y -= 40
                if y < 100:  # new page
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y = height - 50

            c.save()
            buffer.seek(0)
            return buffer

        # Download buttons
        st.download_button("📥 Download as Word", data=save_as_docx(mcqs),
                           file_name="mcqs.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.download_button("📥 Download as PDF", data=save_as_pdf(mcqs),
                           file_name="mcqs.pdf", mime="application/pdf")
